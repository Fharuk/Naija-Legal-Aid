from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import datetime

class LegalDocBuilder:
    """
    Generates professionally formatted .docx files for legal correspondence.
    """
    
    @staticmethod
    def generate_letter(user_name: str, letter_data: dict) -> io.BytesIO:
        doc = Document()
        
        # Styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # 1. Header (Date)
        today = datetime.date.today().strftime("%B %d, %Y")
        p = doc.add_paragraph(today)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph("") # Spacing

        # 2. Recipient Block
        recipient = letter_data.get('recipient_type', 'RECIPIENT').upper()
        doc.add_paragraph(f"TO: THE {recipient}")
        doc.add_paragraph("[Address - Please Fill Manually]")
        doc.add_paragraph("[City, State]")
        
        doc.add_paragraph("") 
        
        # 3. Salutation
        doc.add_paragraph("Dear Sir/Madam,")
        doc.add_paragraph("") 
        
        # 4. Subject Line (Bold & Centered/Uppercase)
        subject_text = f"FORMAL NOTICE: REGARDING LEGAL RIGHTS AND OBLIGATIONS"
        p_sub = doc.add_paragraph()
        run = p_sub.add_run(subject_text)
        run.bold = True
        run.underline = True
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("") 
        
        # 5. Body Text
        body_text = letter_data.get('formal_body', 'Content missing.')
        p_body = doc.add_paragraph(body_text)
        p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # 6. Closing Block
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("Yours faithfully,")
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("__________________________")
        doc.add_paragraph(user_name)
        
        # Output
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer